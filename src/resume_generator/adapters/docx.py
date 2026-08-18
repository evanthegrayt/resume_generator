import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(90, 90, 90)
BLACK = RGBColor(20, 20, 20)


def set_font(run, size=9.2, bold=False, italic=False, color=BLACK):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_para_spacing(p, before=0, after=0, line=1.02):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_rule(paragraph, color="D9E2EF"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=6, after=2, line=1.0)
    add_rule(p)
    r = p.add_run(text.upper())
    set_font(r, size=10.5, bold=True, color=BLUE)
    return p


def add_text_line(doc, parts, before=0, after=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, before=before, after=after, line=1.02)
    for text, opts in parts:
        r = p.add_run(text)
        set_font(
            r,
            size=opts.get("size", 9.2),
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
            color=opts.get("color", BLACK),
        )
    return p


def add_bullet(doc, text, keep_with_next=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.0)
    p.paragraph_format.line_spacing = 1.02
    p.paragraph_format.keep_with_next = keep_with_next
    r = p.add_run(text)
    set_font(r, size=8.75)
    return p


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Normal"].font.size = Pt(9.2)
    return doc


def add_header(doc, resume):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title, before=0, after=0, line=1.0)
    r = title.add_run(resume.contact.name)
    set_font(r, size=18, bold=True, color=BLUE)

    headline = doc.add_paragraph()
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(headline, after=1, line=1.0)
    r = headline.add_run(resume.variant.headline)
    set_font(r, size=10.2, bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(contact, after=4, line=1.0)
    r = contact.add_run(" | ".join(resume.contact.display_items()))
    set_font(r, size=8.4, color=GRAY)


def add_summary(doc, resume):
    add_heading(doc, "Summary")
    add_text_line(doc, [(resume.variant.summary, {"size": 8.9})], after=1)


def add_skills(doc, resume):
    add_heading(doc, "Technical Skills")
    for skill in resume.skills:
        add_text_line(
            doc,
            [
                (f"{skill.label}: ", {"bold": True, "size": 8.55}),
                (skill.value, {"size": 8.55}),
            ],
            after=0,
        )


def add_experience(doc, resume):
    add_heading(doc, "Experience")
    for company in resume.experience:
        company_p = add_text_line(
            doc,
            [
                (company.name, {"bold": True, "size": 9.5}),
                (f" | {company.location}", {"bold": True, "size": 9.5}),
            ],
            before=2.5,
            after=0,
        )
        company_p.paragraph_format.keep_with_next = True

        for role_index, role in enumerate(company.roles):
            title_p = add_text_line(
                doc,
                [
                    (role.title, {"italic": True, "size": 9.2}),
                    (f" | {role.dates}", {"italic": True, "size": 9.2, "color": GRAY}),
                ],
                before=0 if role_index == 0 else 2,
                after=0,
            )
            title_p.paragraph_format.keep_with_next = True
            for index, bullet in enumerate(role.bullets):
                add_bullet(doc, bullet, keep_with_next=index == 0 and len(role.bullets) > 1)


def add_open_source(doc, resume):
    add_heading(doc, "Open Source")
    for bullet in resume.open_source:
        add_bullet(doc, bullet)


def add_education(doc, resume):
    add_heading(doc, "Education")
    add_text_line(
        doc,
        [
            (resume.education.school, {"bold": True, "size": 8.8}),
            (f" | {resume.education.location} | {resume.education.details}", {"size": 8.8}),
        ],
        after=0,
    )


def build_docx(resume, out_dir):
    doc = configure_document()
    add_header(doc, resume)
    add_summary(doc, resume)
    add_skills(doc, resume)
    add_experience(doc, resume)
    add_open_source(doc, resume)
    add_education(doc, resume)

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{resume.variant.stem}.docx"
    doc.save(out_path)
    return out_path


def soffice_path():
    configured = os.environ.get("SOFFICE")
    if configured:
        return configured
    return shutil.which("soffice") or shutil.which("libreoffice")


def export_pdf(docx_path):
    soffice = soffice_path()
    if not soffice:
        raise RuntimeError("LibreOffice/soffice was not found; rerun with --no-pdf to skip PDF export.")

    with tempfile.TemporaryDirectory(prefix="resume-soffice-") as tmpdir:
        profile = Path(tmpdir) / "profile"
        cmd = [
            soffice,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ]
        subprocess.run(cmd, check=True, capture_output=True, text=True)
    return docx_path.with_suffix(".pdf")
